from flask import Flask, render_template, request, send_file
from docx import Document
import os
import html
import tempfile
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.text import MIMEText
import smtplib
from docx2pdf import convert

app = Flask(__name__)

def replace_in_table(table, form_data):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for key, value in form_data.items():
                    if f'{{{key}}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{key}}}', value)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    try:
        # Sanitize user input
        form_data = {
            key: html.escape(request.form[key])
            for key in request.form
        }

        # Load the DOCX template
        template_path = 'template.docx'
        doc = Document(template_path)

        # Replace placeholders in paragraphs and tables
        for paragraph in doc.paragraphs:
            for key, value in form_data.items():
                if f'{{{key}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{key}}}', value)
        for table in doc.tables:
            replace_in_table(table, form_data)

        # Replace 'Modalidade' placeholder
        for paragraph in doc.paragraphs:
            if '{Modalidade}' in paragraph.text:
                modalidade_values = request.form.getlist('modalidade')
                modalidade_text = ', '.join(html.escape(value) for value in modalidade_values)
                paragraph.text = paragraph.text.replace('{Modalidade}', modalidade_text)

        # Replace 'ResidueTest' placeholder
        for paragraph in doc.paragraphs:
            if '{ResidueTest}' in paragraph.text:
                residue_test_values = request.form.getlist('residueTest')
                residue_test_text = ', '.join(html.escape(value) for value in residue_test_values)
                paragraph.text = paragraph.text.replace('{ResidueTest}', residue_test_text)

        # Create temporary DOCX file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_doc:
            doc.save(temp_doc.name)
            temp_doc_path = temp_doc.name

        # Convert DOCX to PDF using python-docx2pdf
        pdf_path = temp_doc_path.replace(".docx", ".pdf")
        convert(temp_doc_path, pdf_path)

        # Remove the temporary DOCX file
        os.remove(temp_doc_path)

        # Send the PDF via email
        send_email_with_attachment(pdf_path, request.form.get('recipient_email'))

        # Return the PDF for download
        return send_file(pdf_path, as_attachment=True, download_name='filled_form.pdf')

    except Exception as e:
        # Log the exception for debugging
        app.logger.error(f"An error occurred: {str(e)}")
        return f"An error occurred: {str(e)}"

def send_email_with_attachment(pdf_path, recipient_email):
    # Email credentials
    sender_email = 'prescricoesineb@gmail.com'
    sender_password = 'affr ebla htup lznh'
    subject = 'Here is your PDF'
    body = 'Please find the attached PDF document.'

    # Create email
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    # Attach PDF file
    with open(pdf_path, 'rb') as attachment:
        part = MIMEApplication(attachment.read(), _subtype="pdf")
        part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(pdf_path))
        msg.attach(part)

    # Send email
    with smtplib.SMTP('smtp.gmail.com', 587) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, recipient_email, msg.as_string())

if __name__ == '__main__':
    app.run(debug=True)
